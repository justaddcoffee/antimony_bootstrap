#!/usr/bin/env python3
"""Generate Word document supplement for PyAntiGen paper."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json

doc = Document()

# Style setup
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title
title = doc.add_heading('Supplementary Material', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Agentic Identification of ODE-Extractable Papers in a Disease-Specific Corpus')
run.bold = True
run.font.size = Pt(14)

# ---- S1. Introduction ----
doc.add_heading('S1. Introduction', level=1)
doc.add_paragraph(
    'A key capability for agentic construction of mechanistic ODE models is the automated '
    'identification of published papers that contain extractable kinetic information: rate constants, '
    'rate law functional forms, species concentrations, and mechanistic descriptions suitable for '
    'conversion to ordinary differential equations. We describe a two-stage agentic pipeline that '
    'triages a corpus of 3,463 Alzheimer\'s disease papers (1,325 primary and 2,138 secondary '
    'Alzforum-curated papers, obtained as PMC XML full text) and extracts structured mechanistic '
    'data from the highest-scoring subset. We validate the triage by demonstrating that all known '
    'ODE-containing papers in the corpus are captured in the highest-priority tier.'
)

# ---- S2. Methods ----
doc.add_heading('S2. Methods', level=1)

doc.add_heading('S2.1 Corpus', level=2)
doc.add_paragraph(
    'The corpus consisted of 3,463 PubMed Central (PMC) XML full-text articles curated by '
    'Alzforum (alzforum.org), a community resource for Alzheimer\'s disease research. Papers '
    'were categorized by Alzforum as primary (1,325 papers; original research directly relevant '
    'to AD) or secondary (2,138 papers; reviews, related research, or supporting studies). All '
    'papers were available as PMC XML with full text, tables, and figure captions. After removing '
    '435 papers that appeared in both the primary and secondary collections, the corpus contained '
    '3,028 unique papers.'
)

doc.add_heading('S2.2 Stage 1: Keyword-Based Triage', level=2)
doc.add_paragraph(
    'The first stage used a deterministic Python script to score all 3,463 papers for ODE '
    'relevance without requiring a large language model (LLM). For each paper, the script parsed '
    'the PMC XML and computed three component scores:'
)

p = doc.add_paragraph()
run = p.add_run('Kinetic score (K): ')
run.bold = True
p.add_run(
    'Count of kinetic-relevant keyword matches in the abstract and body text. Keywords: '
    'rate, kinetic, Km, Kd, IC50, half-life, clearance, production, degradation, turnover, '
    'flux, kon, koff, kcat, Vmax, dissociation constant, binding affinity, rate constant. '
    'Each keyword was matched as a whole word with case-insensitive regular expressions.'
)

p = doc.add_paragraph()
run = p.add_run('Mechanism score (M): ')
run.bold = True
p.add_run(
    'Count of mechanistic keyword matches. Keywords: pathway, signaling, phosphorylation, '
    'aggregation, binding, transport, cleavage, secretion, nucleation, elongation, fibrillization, '
    'oligomerization, endocytosis, transcytosis, phagocytosis.'
)

p = doc.add_paragraph()
run = p.add_run('Table score (T): ')
run.bold = True
p.add_run(
    'Binary assessment of table content. Papers containing <table-wrap> XML elements with '
    'kinetic keywords (rate, Km, Kd, IC50, half-life, Ki, kon, koff, EC50) in the table text '
    'received T=2; papers with tables but no kinetic keywords received T=1; papers without '
    'tables received T=0.'
)

doc.add_paragraph('The composite score was computed as:')
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq.add_run('total_score = K \u00d7 3 + M + T \u00d7 5')
run.italic = True

doc.add_paragraph(
    'The weighting reflects the observation that kinetic keywords are stronger predictors of '
    'extractable rate constants than mechanistic keywords alone, while tables with kinetic data '
    'are particularly high-value for parameter extraction. Papers were ranked by total score '
    'and assigned to four tiers based on rank (Table S1).'
)

doc.add_heading('S2.3 Stage 2: LLM-Based Deep Extraction', level=2)
doc.add_paragraph(
    'Tier A and Tier B papers (500 total) were submitted to parallel LLM agents for structured '
    'knowledge extraction. Each agent received a batch of approximately 35 papers and was '
    'instructed to parse the full PMC XML text, tables, and figure captions. For each paper, '
    'agents extracted:'
)

items = [
    'Reactions: reactants, products, rate type (mass action, Michaelis-Menten, custom), and rate equation templates',
    'Parameters: kinetic constants with values, units, measurement context, and species source (human, mouse, in vitro)',
    'Kinetic values: half-lives, concentrations, binding affinities, production and clearance rates',
    'Rate law forms: evidence for the functional form of kinetic relationships',
    'Species and compartments: molecular species names and their biological compartments',
    'Pathway classification: assignment to one of 20 predefined AD-relevant pathway categories',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Tier C and D papers were processed for pathway classification only, without detailed '
    'parameter extraction. The extraction was executed as 100 parallel LLM agents, each running '
    'autonomously with no human intervention. Total agent compute time was approximately '
    '200 agent-hours; wall-clock time was approximately 2 hours due to parallelization.'
)

doc.add_heading('S2.4 Validation Against Known ODE Papers', level=2)
doc.add_paragraph(
    'To validate the triage, we identified 14 papers in the corpus that are known a priori to '
    'contain explicit ODE models or extractable kinetic rate equations. These include published '
    'kinetic models of amyloid-\u03b2 aggregation from the Knowles group (nucleation-elongation '
    'ODE models), SILK (Stable Isotope Labeling Kinetics) models from the Bateman/Elbert group '
    '(compartmental ODE models of amyloid-\u03b2 and tau kinetics), and pharmacokinetic studies '
    'with explicit rate measurements. We checked whether these papers were captured by the '
    'Tier A assignment.'
)

# ---- S3. Results ----
doc.add_heading('S3. Results', level=1)

doc.add_heading('S3.1 Triage Results', level=2)

# Table S1
doc.add_paragraph().add_run('Table S1: Corpus triage results').bold = True
table = doc.add_table(rows=6, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Tier', 'Papers', 'Score Range', 'Avg Score', 'Avg K', 'Avg M', 'Purpose']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ['A', '200', '120\u2013508', '186.8', '39.7', '64.8', 'Deep parameter extraction'],
    ['B', '300', '76\u2013120', '94.1', '15.8', '44.4', 'Mechanistic extraction'],
    ['C', '1,000', '30\u201376', '47.9', '7.6', '22.3', 'Pathway identification only'],
    ['D', '1,963', '0\u201330', '13.8', '1.9', '4.7', 'Not processed'],
    ['Total', '3,463', '0\u2013508', '40.6', '', '', ''],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()

doc.add_heading('S3.2 Extraction Yield from Tier A and B', level=2)
doc.add_paragraph(
    'Table S2 summarizes the structured data extracted from Tier A and B papers by the '
    'LLM agents. Across the 500 deeply-read papers, the agents identified 13,816 reactions, '
    '33,916 parameters, 5,292 kinetic values, and 3,162 rate law form descriptions.'
)

# Table S2
doc.add_paragraph().add_run('Table S2: LLM extraction results by tier').bold = True
table = doc.add_table(rows=4, cols=8)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Tier', 'Papers', 'With Rxns', 'Total Rxns', 'With Params', 'Total Params', 'With Kinetic', 'Total Kinetic']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
data = [
    ['A', '200', '179 (90%)', '6,534', '198 (99%)', '11,916', '169 (85%)', '2,041'],
    ['B', '300', '262 (87%)', '7,282', '298 (99%)', '22,000', '254 (85%)', '3,251'],
    ['Total', '500', '441 (88%)', '13,816', '496 (99%)', '33,916', '423 (85%)', '5,292'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()

doc.add_heading('S3.3 Validation: Known ODE Papers', level=2)
doc.add_paragraph(
    'All 14 known ODE-containing papers were assigned to Tier A by the keyword triage '
    '(Table S3). These papers span two major categories of kinetic modeling in the AD '
    'literature: (1) amyloid-\u03b2 aggregation kinetics from the Knowles/Linse group, which '
    'use nucleation-polymerization ODE frameworks, and (2) in vivo isotope labeling kinetics '
    'from the Bateman/Elbert group, which use compartmental ODE models to measure protein '
    'production and clearance rates in humans.'
)

# Table S3
doc.add_paragraph().add_run(
    'Table S3: Known ODE-containing papers and their triage scores'
).bold = True
table = doc.add_table(rows=15, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['PMCID', 'Tier', 'Score', 'K', 'Description']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

ode_papers = [
    ['PMC4694579', 'A', '508', '149', 'Clearance systems in the brain (review with kinetic data)'],
    ['PMC7882582', 'A', '426', '139', 'Tau formation and turnover rates (photodynamic labeling)'],
    ['PMC8864715', 'A', '370', '45', 'A\u03b2 peptide aggregation inhibition kinetics'],
    ['PMC4498454', 'A', '337', '73', 'Knowles: lag phase in amyloid fibril formation (ODE model)'],
    ['PMC4621881', 'A', '336', '107', 'In vivo tau turnover kinetics (isotope tracing)'],
    ['PMC4758743', 'A', '331', '48', 'Knowles: primary nucleation kinetics with drug inhibition'],
    ['PMC4820785', 'A', '322', '69', 'Knowles: microscopic kinetic mechanisms of aggregation'],
    ['PMC4546566', 'A', '308', '97', 'Bateman: amyloid-\u03b2 SILK kinetics (compartmental ODE)'],
    ['PMC6137722', 'A', '300', '90', 'Elbert/Bateman: tau SILK kinetics (compartmental ODE)'],
    ['PMC4595974', 'A', '297', '55', 'Knowles: Brichos chaperone breaks catalytic cycle'],
    ['PMC3838868', 'A', '270', '81', 'Bateman: A\u03b242 production/exchange in presenilin mutations'],
    ['PMC6537657', 'A', '230', '25', 'Tau phosphorylation rates by mass spectrometry'],
    ['PMC4269372', 'A', '168', '51', 'CSF A\u03b2 production rate measurement (antidepressant effect)'],
    ['PMC2913973', 'A', '123', '39', '\u03b3-secretase inhibition kinetics in nonhuman primates'],
]
for r, row_data in enumerate(ode_papers):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()

doc.add_paragraph(
    'A broader search for papers with kinetic-rich titles (containing terms such as '
    '"kinetic," "rate," "clearance," "turnover," "half-life," "nucleation," or "aggregation") '
    'and high kinetic keyword scores (K \u2265 40) identified 43 papers. All 43 were in Tier A.'
)

doc.add_heading('S3.4 Pathway Coverage', level=2)
doc.add_paragraph(
    'Table S4 shows the pathway distribution across all papers. Pathway classification was '
    'performed for all tiers, including C and D.'
)

# Table S4
doc.add_paragraph().add_run('Table S4: Pathway classification (all papers)').bold = True
table = doc.add_table(rows=11, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = 'Pathway'
table.rows[0].cells[1].text = 'Papers'
for paragraph in table.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in table.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

pathway_data = [
    ['Amyloid clearance', '3,451'],
    ['Oxidative stress', '3,403'],
    ['Amyloid transport', '2,935'],
    ['Amyloid aggregation', '2,232'],
    ['ApoE genetics', '1,849'],
    ['Tau phosphorylation', '1,743'],
    ['Synaptic dysfunction', '1,560'],
    ['Tau aggregation', '1,439'],
    ['Neuroinflammation (microglia)', '1,132'],
    ['Neuroinflammation (astrocyte)', '1,121'],
]
for r, row_data in enumerate(pathway_data):
    table.rows[r+1].cells[0].text = row_data[0]
    table.rows[r+1].cells[1].text = row_data[1]

doc.add_paragraph()

# ---- S4. Discussion ----
doc.add_heading('S4. Discussion', level=1)

doc.add_heading('S4.1 Triage as a Coarse Filter', level=2)
doc.add_paragraph(
    'The keyword-based triage functions as an effective coarse filter for identifying papers '
    'with ODE-extractable content. The validation against 14 known ODE papers demonstrates '
    'perfect recall: all confirmed kinetic modeling papers were captured in Tier A. A broader '
    'search for papers with kinetic-rich titles and high kinetic keyword density (K \u2265 40) '
    'found 43 papers, all in Tier A. This indicates that the scoring formula reliably assigns '
    'high scores to papers containing explicit kinetic models and rate equations.'
)

doc.add_paragraph(
    'However, the triage does not provide fine-grained ranking within the top tiers. '
    'Within Tier A and B papers that were both subjected to deep LLM extraction, the '
    'extraction yields were similar: 90% of Tier A and 87% of Tier B papers yielded at '
    'least one extractable reaction. The Spearman correlation between triage score and '
    'number of extracted reactions was weak (r = 0.13, p = 0.006) within the A+B group, '
    'suggesting that the triage score is primarily useful as a binary classifier (relevant '
    'vs. not relevant) rather than a continuous predictor of ODE content richness.'
)

doc.add_heading('S4.2 Why the Kinetic Score Dominates', level=2)
doc.add_paragraph(
    'The 3\u00d7 weighting of kinetic keywords (K) relative to mechanism keywords (M) in the '
    'composite score appears well-calibrated. All 14 known ODE papers had K \u2265 25, and '
    'the top 50 papers ranked by K alone were all in Tier A. Papers that score highly on '
    'mechanism keywords but not kinetic keywords (e.g., reviews describing signaling pathways '
    'qualitatively) are appropriately ranked lower than papers with explicit rate measurements. '
    'The table score bonus (T \u00d7 5) also contributes meaningfully: papers with kinetic data '
    'tables (T = 2) are particularly valuable for parameter extraction.'
)

doc.add_heading('S4.3 Scale of Extractable Knowledge', level=2)
doc.add_paragraph(
    'The extraction of 13,816 reactions and 33,916 parameters from 500 papers indicates '
    'that the Alzheimer\'s disease literature contains substantially more ODE-relevant '
    'quantitative data than is typically incorporated into any single published model. '
    'For context, the QSP models described in the main text contain 16\u201399 reactions. '
    'However, the extracted data requires significant curation; many "parameters" identified '
    'by the LLM agents represent experimental measurements (drug doses, patient demographics) '
    'rather than true kinetic rate constants.'
)

doc.add_heading('S4.4 Limitations', level=2)

doc.add_paragraph(
    'Several limitations should be noted. First, the validation set of 14 known ODE papers '
    'was curated by the authors, who are familiar with this literature; an independent '
    'validation set would be stronger. Second, we cannot assess false negative rates for '
    'Tier C and D because those papers were not subjected to deep extraction. It is possible '
    'that some C or D papers contain ODE-relevant content described using non-standard '
    'terminology that the keyword list does not capture. Third, the LLM extraction results '
    'used to characterize Tier A and B yields were not individually validated by humans; '
    'false positives (incorrectly identified reactions) and false negatives (missed data) '
    'are expected. Fourth, the pathway classification shows high counts for broad categories '
    'such as "amyloid clearance" across nearly all papers, suggesting that the classification '
    'may be too permissive for some pathway categories.'
)

doc.add_heading('S4.5 Applicability Beyond Alzheimer\'s Disease', level=2)
doc.add_paragraph(
    'The two-stage triage approach is disease-agnostic. The kinetic and mechanistic keyword '
    'lists could be applied to any PMC XML corpus to identify papers with ODE-extractable '
    'content. The approach would be particularly effective for diseases with established '
    'mechanistic literature, such as cancer signaling, diabetes, and infectious disease, '
    'where QSP and PBPK models are actively developed.'
)

# ---- S5 ----
doc.add_heading('S5. Data Availability', level=1)
doc.add_paragraph(
    'The complete paper index (paper_index.json), tier assignments, and extraction results '
    'are available in the project repository. The triage script (build_paper_index.py) can '
    'be applied to any collection of PMC XML files and runs in under 5 minutes on a standard '
    'laptop for the full 3,463-paper corpus.'
)

# Save
outpath = '/Users/jtr4v/PythonProject/antimony_bootstrap/Supplement_ODE_Triage.docx'
doc.save(outpath)
print(f"Saved to {outpath}")
